###############################################################
##                                                           ##
##  scoreAI                                                  ##
##  -------                                                  ##
##                                                           ##
##  Python script for processing transcribed and scored      ##
##  autobiographical memory narratives.                      ##
##                                                           ##
##  Current public version: build 10[20200511]               ##
##                                                           ##
##                                                           ##
##  If you use this software, please cite:                   ##
##  * Wardell, V., Esposito, C. L., Madan, C. R., & Palombo, ##
##      D. J. (preprint). Semi-automated transcription and   ##
##      scoring of autobiographical memory narratives.       ##
##                                                           ##
##  For documentation and further information, see:          ##
##     https://github.com/cMadan/scoreAI                     ##
##                                                           ##
###############################################################


###############################################################
## SECTION 1: Config
# folder with memory docs
pathIn = 'input'
pathOut = 'output'

# how many memories are we expecting per document?
nMemories = 3


###############################################################
## SECTION 2: Load dependencies
from docx import Document #python-docx package
from os import listdir
from os.path import isfile, join
import pandas as pd
from datetime import date
from warnings import warn


###############################################################
## SECTION 3: Expert settings
# memory scoring codes
tags = ['Int_EV','Int_PERC','Int_EMO','Int_PL','Int_TM','Ext_EV','Ext_SEM','Ext_REP','Ext_OTH']

# file lists
docs = [f for f in listdir(pathIn) if isfile(join(pathIn, f))]

# min length (characters) of a valid memory response,
# used to parse doc and skip empty runs
# only used by getResponse, but not getPara
minresponselength = 25


###############################################################
## SECTION 4: Define re-usable functions
# seek next paragraph that includes specific text string
def seekPara(para,string):
    search = True
    while search:
        para += 1
        try:
            # strip removes trailing spaces
            #text = d.paragraphs[para].runs[0].text.strip()
            # above approach is too brittle to idiosyncrasies of Word XML formatting
            text = getPara(para).strip()
            # check if line matches specific search string
            compare = text == string
            if compare:
                search = False # found it!
            elif para == len(runcount)-1:
                # didn't find the string,
                # shouldn't ever happen
                para = False
                search = False # abort
        except:
            # do nothing, will fail often
            False
    return para

# pull paragraph text
def getPara(para):
    response = ''
    for run in d.paragraphs[para].runs:
        response += run.text
    response = response.strip()
    return response

# pull paragraph text for a transcribed response
def getResponse(para):
    search = True
    while search:
       minlength = runcount[para][1]>minresponselength
       if minlength:
           search = False # found it!
       else:
           para += 1
    # para found, now compile text across runs
    response = getPara(para)
    return response

# get counts for each tag
def countTag(response):
    tagCount = []
    for tag in tags:
        count = response.count(tag)
        tagCount.append(count)
    return tagCount

# go through full document and getting ER values
def getER():
    # get episodic richness values for each memory
    search = True
    ERlist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()
        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:3] == '[ER':
            # get the specific ER values
            ERlist.append(text[4])
    if len(ERlist) != nMemories:
        warn('WARNING: In %s, number of episodic richness codes does not match number of memories expected. %g ER codes found, expected %g.' % (doc,len(ERlist),nMemories))
        print(ERlist)
    return ERlist


###############################################################
## SECTION 5: Process the data
# cycle through each doc in input folder
for doc in docs:
    print('Processing '+doc)

    # load document
    d = Document(join(pathIn,doc))

    ## parse document
    # enumerate paragraph runs
    runcount = [len(d.paragraphs[p].runs) for p in range(len(d.paragraphs))]
    textcount = []
    for p in d.paragraphs:
        text = ''
        for runs in p.runs:
            text += runs.text
        textcount.append(len(text))
    runcount = list(zip(runcount,textcount))

    # get the Participant ID
    cover = d.tables[0].rows[0].cells[0].text
    cover = cover.split('\n')
    subID = cover[1].replace('Participant ID: ','').strip()

    # get episodic richness values
    ERlist = getER()

    # get start and end paras for each memory
    para = 0
    paraM = []
    for M in list(range(1,nMemories+1)):
        # find para number for Memory M
        para = seekPara(para,'Memory '+str(M))
        paraM.append(para)
        print(str(M)+'...', end = '')
    if len(paraM) != nMemories:
        warn('WARNING: In %s, Number of memories found does not match number of memories expected. %g memories found, expected %g.' % (doc,len(paraM),nMemories))

    # add last run number as end value for para intervals
    paraM.append(len(runcount))

    # get text and counts for each memory
    responseCounts = []
    for M in list(range(1,nMemories+1)):
        response = ''
        for para in list(range(paraM[M-1],paraM[M])):
            response += getPara(para)
            response += '|'
        counts = countTag(response)
        responseCounts.append(counts)

    ## output data
    # init dict first
    data_sub = {
        'ParticipantID': [subID] * nMemories, # repeat subID several times, like repmat
        'Memory': list(range(1,nMemories+1))
    }
    # restructure tags to have list of each tag and count for each memory
    counts = dict(zip(tags, list(map(list,zip(*responseCounts))) ))
    # concatenate dicts, to merge in the tag counts
    data_sub.update(counts)
    # concat in the episodic richness values
    data_sub.update({'ER':ERlist})
    # merge data to common record, or create it if not exists
    try:
        data_all = { key:data_all.get(key,[])+data_sub.get(key,[]) for key in data_all.keys() }
    except NameError:
        data_all = data_sub
    print('done.')

###############################################################
## SECTION 6: Output the counts
# convert data dict into a dataframe
data_all = pd.DataFrame(data_all)
# write to file
fOut = 'scoring_%s_n%g.csv' % (date.today().strftime('%Y%m%d'),len(docs))
data_all.to_csv(join(pathOut,fOut),index=False)


###############################################################
## fin
